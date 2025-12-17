#!/usr/bin/env python3
"""
Create Evaluation Package from Data

Main automation script that:
1. Assembles markdown from various sources (Coda, LaTeX, PDFs)
2. Creates PubPub package structure
3. Imports content (using Word documents for proper table rendering)
4. Sets up metadata, connections, and attributions

This is the general-purpose script for any evaluation package.
"""

import sys
import json
import os
import tempfile
from pathlib import Path
from typing import Dict, List, Optional

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from pypubpub import Pubshelper_v6
from pypubpub.Pubv6 import EvaluationPackage

from package_assembler import PackageAssembler, PaperMetadata, EvaluationData, EvaluationPackageData
from template_generator import TemplateGenerator


class EvaluationPackageCreator:
    """Creates complete evaluation packages in PubPub."""

    def __init__(
        self,
        email: str,
        password: str,
        community_url: str,
        community_id: str
    ):
        """
        Initialize creator with PubPub credentials.

        Args:
            email: PubPub login email
            password: PubPub password
            community_url: Community URL (e.g., 'https://unjournal.pubpub.org')
            community_id: Community UUID
        """
        self.pubhelper = Pubshelper_v6(
            community_url=community_url,
            community_id=community_id,
            email=email,
            password=password
        )
        self.pubhelper.login()
        self.assembler = PackageAssembler()
        self.template_generator = TemplateGenerator()

    def create_package(
        self,
        package_data: EvaluationPackageData,
        draft_mode: bool = True,
        output_dir: Optional[Path] = None
    ) -> Dict:
        """
        Create complete evaluation package in PubPub.

        Args:
            package_data: Complete package data
            draft_mode: If True, don't publish evaluator names yet
            output_dir: Optional directory to save markdown files

        Returns:
            Dict with created pub IDs and info
        """
        print(f"\n{'='*60}")
        print(f"Creating Evaluation Package")
        print(f"Paper: {package_data.paper.title}")
        print(f"Evaluations: {len(package_data.evaluations)}")
        print(f"Mode: {'DRAFT (anonymous)' if draft_mode else 'FINAL (with names)'}")
        print(f"{'='*60}\n")

        # Step 1: Create PubPub package structure first (so we have URLs)
        print("Step 1: Creating PubPub package structure...")

        # Prepare evaluation configs for EvaluationPackage
        eval_configs = []
        for i, evaluation in enumerate(package_data.evaluations):
            # In draft mode or if not public, don't add author info yet
            if draft_mode or not evaluation.is_public:
                eval_configs.append({})  # Empty placeholder
            else:
                # In final mode with public evaluator, add attribution
                eval_config = {}
                if evaluation.evaluator_name:
                    eval_config['author'] = {
                        'name': evaluation.evaluator_name
                    }
                    if evaluation.evaluator_orcid:
                        eval_config['author']['orcid'] = evaluation.evaluator_orcid
                    if evaluation.evaluator_affiliation:
                        eval_config['author']['affiliation'] = evaluation.evaluator_affiliation
                eval_configs.append(eval_config)

        # Create package (structure only, no content yet)
        pkg = EvaluationPackage(
            doi=package_data.paper.doi,
            url=package_data.paper.url if not package_data.paper.doi else None,
            evaluation_manager_author=package_data.evaluation_manager,
            evaluations=eval_configs,
            email=self.pubhelper.email,
            password=self.pubhelper.password,
            community_url=self.pubhelper.community_url,
            community_id=self.pubhelper.community_id,
            autorun=True
        )

        print(f"  Created evaluation summary pub: {pkg.eval_summ_pub['id']}")
        for i, (pub_id, pub_obj) in enumerate(pkg.activePubs, 1):
            print(f"  Created evaluation {i} pub: {pub_id}")

        # Step 2: Generate evaluation URLs for hyperlinking
        print("\nStep 2: Generating content with hyperlinked evaluator names...")
        evaluation_links = []
        for pub_id, pub_obj in pkg.activePubs:
            eval_url = f"{self.pubhelper.community_url}/pub/{pub_obj['slug']}"
            evaluation_links.append(eval_url)
            print(f"  Evaluation URL: {eval_url}")

        # Step 3: Generate HTML content using template generator
        print("\nStep 3: Generating HTML templates...")

        # Build evaluation data for template generator
        template_evaluations = []
        for i, evaluation in enumerate(package_data.evaluations):
            eval_data = {
                'name': evaluation.evaluator_name if (not draft_mode and evaluation.is_public) else f'Evaluator {i+1}',
                'ratings': evaluation.ratings or {},
                'summary': evaluation.summary or '',
            }
            template_evaluations.append(eval_data)

        # Generate summary HTML with hyperlinked evaluator names
        paper_info = {
            'title': package_data.paper.title,
            'authors': package_data.paper.authors,
            'doi': package_data.paper.doi,
        }

        summary_html = self._generate_summary_html(
            paper_info=paper_info,
            evaluations=template_evaluations,
            manager_summary=package_data.manager_summary,
            evaluation_links=evaluation_links
        )
        print(f"  Generated summary HTML ({len(summary_html)} chars)")

        # Step 4: Import content using Word document approach for proper table rendering
        print("\nStep 4: Importing content (using Word format for tables)...")

        # Import summary using Word document approach
        print(f"  Importing summary...")
        self._import_html_via_word(pkg.eval_summ_pub['id'], summary_html, 'evaluation_summary.docx')
        print(f"    Summary content imported")

        # Import individual evaluations
        for i, ((pub_id, pub_obj), evaluation) in enumerate(zip(pkg.activePubs, package_data.evaluations), 1):
            print(f"  Importing evaluation {i}...")
            eval_html = self._generate_evaluation_html(paper_info, evaluation, i, draft_mode)
            self._import_html_via_word(pub_id, eval_html, f'evaluation_{i}.docx')
            print(f"    Evaluation {i} content imported")

        # Step 5: Summary
        print(f"\n{'='*60}")
        print("Package creation complete!")
        print(f"{'='*60}\n")

        print("Created publications:")
        print(f"  Summary: {self.pubhelper.community_url}/pub/{pkg.eval_summ_pub['slug']}")
        for i, (pub_id, pub_obj) in enumerate(pkg.activePubs, 1):
            print(f"  Evaluation {i}: {self.pubhelper.community_url}/pub/{pub_obj['slug']}")

        if draft_mode:
            print("\nDRAFT MODE: Evaluator names not added yet")
            print("   After author response, re-run in final mode to add names")

        # Return package info
        return {
            'summary_pub_id': pkg.eval_summ_pub['id'],
            'summary_slug': pkg.eval_summ_pub['slug'],
            'evaluation_pubs': [
                {'id': pub_id, 'slug': pub_obj['slug']}
                for (pub_id, pub_obj) in pkg.activePubs
            ],
            'evaluation_links': evaluation_links
        }

    def _generate_summary_html(
        self,
        paper_info: Dict,
        evaluations: List[Dict],
        manager_summary: Optional[str],
        evaluation_links: List[str]
    ) -> str:
        """Generate evaluation summary HTML with proper table formatting."""
        # Use template generator to get markdown then convert key sections to HTML
        title = paper_info.get('title', 'Untitled Paper')
        num_evals = len(evaluations)
        plural = 's' if num_evals != 1 else ''

        html = f'''<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"><title>Evaluation Summary</title></head>
<body>
<h1>Abstract</h1>
<p>We organized {num_evals} evaluation{plural} of the paper: "{title}". To read these evaluations, please see the links below.</p>

<h2><strong>Evaluations</strong></h2>
'''
        # Evaluations list with hyperlinks
        for i, (eval_data, link) in enumerate(zip(evaluations, evaluation_links), 1):
            name = eval_data.get('name', f'Evaluator {i}')
            html += f'<p>{i}. <a href="{link}">{name}</a></p>\n'

        # Overall ratings section
        html += '''
<h1><strong>Overall ratings</strong></h1>
<p>We asked evaluators to provide overall assessments as well as ratings for a range of specific criteria.</p>
<p><strong>I. Overall assessment</strong> (See footnote)</p>
<p><strong>II. Journal rank tier, normative rating (0-5):</strong> On a 'scale of journals', what 'quality of journal' <em>should</em> this be published in? Note: 0 = lowest/none, 5 = highest/best.</p>
'''
        # Summary table
        html += self.template_generator._generate_summary_table(evaluations)

        html += f'''
<p>See "<a href="#metrics">Metrics</a>" below for a more detailed breakdown of the evaluators' ratings across several categories. To see these ratings in the context of all Unjournal ratings, with some analysis, see our <a href="{self.template_generator.LINKS['data_presentation']}">data presentation here</a>.</p>
<p><a href="{self.template_generator.LINKS['evaluator_guidelines']}">See here</a> for the current full evaluator guidelines, including further explanation of the requested ratings.</p>

<h2><strong>Evaluation Summaries</strong></h2>
'''
        # Evaluation summaries
        for eval_data in evaluations:
            name = eval_data.get('name', 'Anonymous')
            summary = eval_data.get('summary', '*[Summary to be added]*')
            html += f'<h2>{name}</h2>\n<p>{summary}</p>\n'

        # Metrics section
        html += f'''
<h1><strong>Metrics</strong></h1>
<h2>Ratings</h2>
<p><a href="{self.template_generator.LINKS['quantitative_metrics']}">See here</a> for details on the categories below, and the guidance given to evaluators.</p>
'''
        html += self.template_generator._generate_full_ratings_table(evaluations)

        # Journal ranking tiers
        html += f'''
<h2>Journal ranking tiers</h2>
<p><a href="{self.template_generator.LINKS['journal_tiers']}">See here</a> for more details on these tiers.</p>
'''
        html += self.template_generator._generate_journal_tier_table(evaluations)

        # Claims section
        html += '''
<h1><strong>Claim identification and assessment (summary)</strong></h1>
<p><em>For the full discussions, see the corresponding sections in each linked evaluation.</em></p>
'''
        html += self.template_generator._generate_claims_table(evaluations)

        # References
        html += '''
<h2>References</h2>
<p><em>[References to be added]</em></p>
'''

        # Manager's discussion
        html += '''
<h1><strong>Evaluation manager's discussion</strong></h1>
'''
        if manager_summary:
            html += f'<p>{manager_summary}</p>\n'
        else:
            html += "<p><em>[Manager's discussion to be added]</em></p>\n"

        # Process notes
        html += '''
<h1>Unjournal process notes</h1>
<h2>Why we chose this paper</h2>
<p><em>[To be added]</em></p>
<h2>Evaluation process</h2>
<p><em>[To be added]</em></p>
<h3>COI issues</h3>
<p><em>[To be added]</em></p>
</body>
</html>
'''
        return html

    def _generate_evaluation_html(
        self,
        paper_info: Dict,
        evaluation: EvaluationData,
        eval_number: int,
        draft_mode: bool
    ) -> str:
        """Generate individual evaluation HTML."""
        title = paper_info.get('title', 'Untitled Paper')
        name = evaluation.evaluator_name if (not draft_mode and evaluation.is_public) else f'Evaluator {eval_number}'

        html = f'''<!DOCTYPE html>
<html>
<head><meta charset="UTF-8"><title>Evaluation {eval_number}</title></head>
<body>
<h1>Evaluation {eval_number} of "{title}"</h1>
<p><strong>Evaluator:</strong> {name}</p>
'''
        if evaluation.evaluator_affiliation and not draft_mode and evaluation.is_public:
            html += f'<p><strong>Affiliation:</strong> {evaluation.evaluator_affiliation}</p>\n'
        if evaluation.evaluator_orcid and not draft_mode and evaluation.is_public:
            html += f'<p><strong>ORCID:</strong> <a href="https://orcid.org/{evaluation.evaluator_orcid}">{evaluation.evaluator_orcid}</a></p>\n'

        # Summary
        if evaluation.summary:
            html += f'<h2>Summary</h2>\n<p>{evaluation.summary}</p>\n'

        # Review text
        if evaluation.review_text:
            html += f'<h2>Detailed Review</h2>\n<p>{evaluation.review_text}</p>\n'

        # Ratings table
        if evaluation.ratings:
            html += '<h2>Ratings</h2>\n<table>\n'
            html += '  <tr><th><strong>Category</strong></th><th><strong>Rating</strong></th><th><strong>90% CI</strong></th></tr>\n'

            for key, label in self.template_generator.RATING_CATEGORIES:
                if key in evaluation.ratings:
                    rating_data = evaluation.ratings[key]
                    mid_val = self.template_generator._get_rating_value(rating_data)
                    ci_str = self.template_generator._get_ci_string(rating_data)
                    html += f'  <tr><td><strong>{label}</strong></td><td>{mid_val}</td><td>{ci_str}</td></tr>\n'

            # Journal tiers
            for key, label in [('journal_tier_normative', 'Journal tier (normative)'),
                               ('journal_tier_predictive', 'Journal tier (predictive)')]:
                if key in evaluation.ratings:
                    rating_data = evaluation.ratings[key]
                    mid_val = self.template_generator._get_rating_value(rating_data)
                    ci_str = self.template_generator._get_ci_string(rating_data)
                    html += f'  <tr><td><strong>{label}</strong></td><td>{mid_val}</td><td>{ci_str}</td></tr>\n'

            html += '</table>\n'

        html += '</body>\n</html>'
        return html

    def _import_html_via_word(self, pub_id: str, html_content: str, filename: str):
        """Import HTML content to a pub using Word document conversion for proper table rendering."""
        try:
            from docx import Document
            from bs4 import BeautifulSoup
        except ImportError:
            # Fall back to direct markdown if docx not available
            print("    Warning: python-docx not available, falling back to markdown import")
            self.pubhelper.put_template_doc(
                targetPubId=pub_id,
                template_id=None,
                community_id=self.pubhelper.community_id,
                community_url=self.pubhelper.community_url,
                markdown=html_content
            )
            return

        # Convert HTML to Word document
        soup = BeautifulSoup(html_content, 'html.parser')
        doc = Document()

        def add_text_with_formatting(paragraph, element):
            if element.name is None:
                text = str(element)
                if text.strip():
                    paragraph.add_run(text)
            elif element.name in ['strong', 'b']:
                run = paragraph.add_run(element.get_text())
                run.bold = True
            elif element.name in ['em', 'i']:
                run = paragraph.add_run(element.get_text())
                run.italic = True
            elif element.name == 'a':
                run = paragraph.add_run(element.get_text())
                run.underline = True
            elif element.name == 'br':
                paragraph.add_run('\n')
            else:
                for child in element.children:
                    add_text_with_formatting(paragraph, child)

        body = soup.find('body')
        if not body:
            body = soup

        for element in body.children:
            if element.name is None:
                continue
            if element.name in ['h1', 'h2', 'h3']:
                level = int(element.name[1])
                heading = doc.add_heading(level=level)
                add_text_with_formatting(heading, element)
            elif element.name == 'p':
                para = doc.add_paragraph()
                for child in element.children:
                    add_text_with_formatting(para, child)
            elif element.name == 'table':
                rows = element.find_all('tr')
                if rows:
                    max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
                    table = doc.add_table(rows=len(rows), cols=max_cols)
                    table.style = 'Table Grid'
                    for i, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        for j, cell in enumerate(cells):
                            if j < max_cols:
                                cell_para = table.rows[i].cells[j].paragraphs[0]
                                add_text_with_formatting(cell_para, cell)
                    doc.add_paragraph()

        # Save to temp file and upload
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        try:
            file_url = self.pubhelper.upload_file(tmp_path)
            file_size = os.path.getsize(tmp_path)
            self.pubhelper.import_to_pub(pub_id, file_url, filename, file_size)
        finally:
            os.unlink(tmp_path)

    def create_from_coda(
        self,
        coda_data: Dict,
        paper_metadata: PaperMetadata,
        draft_mode: bool = True,
        output_dir: Optional[Path] = None
    ) -> Dict:
        """Create package from Coda form data."""
        # Parse Coda data into EvaluationPackageData
        evaluations = []
        for eval_row in coda_data.get('evaluations', []):
            evaluation = EvaluationData(
                ratings=eval_row.get('ratings', {}),
                review_text=eval_row.get('written_evaluation'),
                evaluator_name=eval_row.get('evaluator_name'),
                evaluator_affiliation=eval_row.get('affiliation'),
                evaluator_orcid=eval_row.get('orcid'),
                is_public=eval_row.get('identified', False),
                comments=eval_row.get('additional_comments')
            )
            evaluations.append(evaluation)

        package_data = EvaluationPackageData(
            paper=paper_metadata,
            evaluations=evaluations,
            manager_summary=coda_data.get('manager_summary'),
            evaluation_manager=coda_data.get('manager_id')
        )

        return self.create_package(package_data, draft_mode, output_dir)

    def create_from_files(
        self,
        paper_metadata: PaperMetadata,
        evaluation_files: List[Dict],
        manager_summary: Optional[str] = None,
        manager_id: Optional[str] = None,
        draft_mode: bool = True,
        output_dir: Optional[Path] = None
    ) -> Dict:
        """Create package from local files."""
        evaluations = []

        for eval_file_info in evaluation_files:
            # Load ratings
            ratings = eval_file_info['ratings']
            if isinstance(ratings, (str, Path)):
                with open(ratings, 'r') as f:
                    ratings = json.load(f)

            # Determine review source
            review_path = eval_file_info.get('review')
            review_type = None
            if review_path:
                review_path = Path(review_path)
                if review_path.suffix == '.tex':
                    review_type = 'latex'
                elif review_path.suffix in ['.doc', '.docx']:
                    review_type = 'word'
                elif review_path.suffix == '.md':
                    review_type = 'markdown'
                else:
                    review_type = 'text'

            evaluation = EvaluationData(
                ratings=ratings,
                evaluator_name=eval_file_info.get('evaluator_name'),
                evaluator_affiliation=eval_file_info.get('affiliation'),
                evaluator_orcid=eval_file_info.get('orcid'),
                is_public=eval_file_info.get('is_public', False),
                review_source_type=review_type,
                review_source_path=review_path
            )
            evaluations.append(evaluation)

        package_data = EvaluationPackageData(
            paper=paper_metadata,
            evaluations=evaluations,
            manager_summary=manager_summary,
            evaluation_manager=manager_id
        )

        return self.create_package(package_data, draft_mode, output_dir)


def main():
    """Example usage and CLI."""
    import argparse

    parser = argparse.ArgumentParser(description='Create evaluation package in PubPub')
    parser.add_argument('--config', required=True, help='Path to JSON config file')
    parser.add_argument('--draft', action='store_true', help='Draft mode (no evaluator names)')
    parser.add_argument('--output-dir', help='Directory to save markdown files')

    args = parser.parse_args()

    # Load configuration
    with open(args.config, 'r') as f:
        config = json.load(f)

    # Create package creator
    creator = EvaluationPackageCreator(
        email=config['pubpub']['email'],
        password=config['pubpub']['password'],
        community_url=config['pubpub']['community_url'],
        community_id=config['pubpub']['community_id']
    )

    # Parse paper metadata
    paper = PaperMetadata(**config['paper'])

    # Create package from evaluation files
    result = creator.create_from_files(
        paper_metadata=paper,
        evaluation_files=config['evaluations'],
        manager_summary=config.get('manager_summary'),
        manager_id=config.get('manager_id'),
        draft_mode=args.draft,
        output_dir=Path(args.output_dir) if args.output_dir else None
    )

    print("\n✓ Package created successfully!")
    print(f"Summary: {result['summary_slug']}")


if __name__ == '__main__':
    main()
